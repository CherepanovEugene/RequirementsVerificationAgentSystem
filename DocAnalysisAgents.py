# -*- coding: utf-8 -*-
# ============================================================
# Мультидокументный анализ требований с эффективным поиском
# ============================================================

import os
import re
import json
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import List, Dict, Optional, Tuple

from dotenv import load_dotenv
from pydantic import BaseModel, Field

# Тексты и документы
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

# LLM и эмбеддинги (GigaChat)
from langchain_gigachat import GigaChat, GigaChatEmbeddings
from langchain_community.vectorstores import FAISS

# LangGraph
from langgraph.graph import StateGraph, END

# Парсинг файлов
import fitz  # PyMuPDF для PDF
try:
    import docx2txt  # для DOCX (опционально)
    HAS_DOCX2TXT = True
except Exception:
    HAS_DOCX2TXT = False

try:
    from docx import Document as DocxDocument  # fallback вариант
    HAS_PYTHON_DOCX = True
except Exception:
    HAS_PYTHON_DOCX = False

# PDF-отчёты
from fpdf import FPDF, XPos, YPos

import zipfile
import re
import time
import random
import threading

# ----------------------------
# Конфиг и логирование
# ----------------------------
load_dotenv()

LOG_DIR = "logs"
os.makedirs(LOG_DIR, exist_ok=True)

def setup_logger(name: str) -> logging.Logger:
    logger = logging.getLogger(name)
    logger.setLevel(logging.INFO)
    fh = logging.FileHandler(os.path.join(LOG_DIR, f"{name}.log"))
    fh.setFormatter(logging.Formatter("%(asctime)s | %(levelname)s | %(message)s"))
    if logger.hasHandlers():
        logger.handlers.clear()
    logger.addHandler(fh)
    return logger

parser_logger = setup_logger("parser_agent")
analysis_logger = setup_logger("analysis_agent")
report_logger = setup_logger("report_agent")
comm_logger = setup_logger("communication_agent")

# ----------------------------
# LLM и эмбеддинги
# ----------------------------
GIGACHAT_API_KEY = os.getenv("GIGACHAT_API_KEY")
llm = GigaChat(
    credentials=GIGACHAT_API_KEY,
    model="GigaChat-Max",
    verify_ssl_certs=False,
    timeout=90,              # увеличенный таймаут
)
embeddings = GigaChatEmbeddings(
    credentials=GIGACHAT_API_KEY,
    model="Embeddings",
    verify_ssl_certs=False
)

# ----------------------------
# Глобальные параметры
# ----------------------------
CHUNK_SIZE = 800
CHUNK_OVERLAP = 120
LONG_PARAGRAPH_WORDS = 350
TOP_K = 8                      # сколько фрагментов брать из векторного поиска
MMR_FETCH_K = 20               # сколько кандидатов рассматривать для MMR
MAX_CTX_CHARS = 5000           # урезаем суммарный контекст для LLM
MAX_WORKERS = 4                # параллелизм по вопросам
MAX_EMBED_TOKENS_START = 480   # стартовый «квази-лимит» по словам
MIN_EMBED_TOKENS = 220         # ниже опускаться не будем
TOKEN_STEP = 40                # шаг уменьшения при ретрае

FONTS_DIR = "fonts"
FONT_REGULAR = os.path.join(FONTS_DIR, "DejaVuSans.ttf")
FONT_BOLD = os.path.join(FONTS_DIR, "DejaVuSans-Bold.ttf")
# при наличии можно добавить курсив:
# FONT_ITALIC = os.path.join(FONTS_DIR, "DejaVuSans-Oblique.ttf")
# Параметры защиты от rate-limit
EMBED_MAX_RETRIES = 6
EMBED_BASE_DELAY = 1.0  # сек
LLM_MAX_RETRIES   = 5
LLM_BASE_DELAY    = 1.0

# Ограничиваем параллельность LLM (чтобы не ловить 429 в чат-эндпойнте)
MAX_WORKERS = 2  # было 4 – снизили для надёжности

def prepare_text_for_pdf(text: str, max_token_len: int = 60) -> str:
    """Нормализует пробелы и вставляет точки переноса в длинные 'слова'."""
    if not text:
        return ""
    t = re.sub(r"\s+", " ", str(text)).strip()

    # Разрешим перенос после некоторых символов
    t = (t.replace("/", "/ ")
           .replace("\\", "\\ ")
           .replace("_", "_ ")
           .replace("-", "- "))

    # Разбиваем слишком длинные непрерывные последовательности
    def _break(m):
        s = m.group(0)
        return " ".join(s[i:i+max_token_len] for i in range(0, len(s), max_token_len))
    t = re.sub(r"\S{" + str(max_token_len) + r",}", _break, t)
    return t

def _token_safe_trim(text: str, max_tokens: int) -> str:
    """Грубая, но быстая аппроксимация: ограничиваем по количеству слов.
    Для кириллицы обычно 1 слово ≈ 1 токен, так что 480 слов ~<514 токенов.
    """
    if not text:
        return text
    words = re.findall(r"\S+", text)
    if len(words) <= max_tokens:
        return text.strip()
    return " ".join(words[:max_tokens]).strip()

def _sanitize_docs_for_embeddings(docs: List[Document], max_tokens: int) -> List[Document]:
    """Возвращает копии документов с обрезанным контентом под лимит max_tokens."""
    sanitized = []
    for d in docs:
        trimmed = _token_safe_trim(d.page_content, max_tokens)
        sanitized.append(Document(page_content=trimmed, metadata=d.metadata))
    return sanitized

def _is_rate_limit_error(exc: Exception) -> bool:
    s = str(exc)
    return ("429" in s) or ("Too Many Requests" in s)

# Семафор для сериализации обращений к Embeddings API
_EMBED_SEMAPHORE = threading.Semaphore(1)

def embed_texts_with_backoff(texts: List[str]) -> List[List[float]]:
    """Безопасная обёртка для embeddings.embed_documents с ретраями и сериализацией."""
    for attempt in range(EMBED_MAX_RETRIES):
        try:
            with _EMBED_SEMAPHORE:
                return embeddings.embed_documents(texts)
        except Exception as e:
            if _is_rate_limit_error(e):
                delay = EMBED_BASE_DELAY * (2 ** attempt) + random.uniform(0, 0.4)
                analysis_logger.warning(f"429 от Embeddings. Повтор через {delay:.1f}s (попытка {attempt+1}/{EMBED_MAX_RETRIES})")
                time.sleep(delay)
                continue
            raise
    # не удалось после ретраев
    raise RuntimeError("Embeddings rate limit after retries")

def embed_query_with_backoff(text: str) -> List[float]:
    return embed_texts_with_backoff([text])[0]

def llm_invoke_with_backoff(prompt: str, timeout: int = 90) -> str:
    """Надёжный вызов chat LLM с ретраями на 429/таймауты."""
    for attempt in range(LLM_MAX_RETRIES):
        try:
            return llm.invoke(prompt, config={"timeout": timeout}).content.strip()
        except Exception as e:
            s = str(e)
            if _is_rate_limit_error(e) or "ReadTimeout" in s or "timeout" in s.lower():
                delay = LLM_BASE_DELAY * (2 ** attempt) + random.uniform(0, 0.6)
                analysis_logger.warning(f"Чат-эндпойнт недоступен ({e}). Повтор через {delay:.1f}s (попытка {attempt+1}/{LLM_MAX_RETRIES})")
                time.sleep(delay)
                continue
            raise
    raise RuntimeError("LLM rate/timeout after retries")

# ----------------------------
# Состояние графа
# ----------------------------
class GraphState(BaseModel):
    file_paths: List[str]                                  # список путей к документам
    questions: List[str]                                   # список вопросов
    docs: List[Document] = Field(default_factory=list)     # фрагменты с метаданными
    vectorstore: Optional[FAISS] = None                    # FAISS-хранилище
    analysis_result: Dict[str, List[Dict]] = Field(default_factory=dict)
    # формат результата по вопросу:
    # {
    #   question: [
    #       {
    #           "answer": str,
    #           "confidence": float,
    #           "citations": [{"id": int, "file": str, "page": int}],
    #           "quotes": [str]
    #       },
    #       ...
    #   ]
    # }

    class Config:
        arbitrary_types_allowed = True

# ----------------------------
# Парсинг документов
# ----------------------------
def _extract_pdf(path: str) -> List[Tuple[str, int]]:
    """Возвращает список (текст_страницы, номер_страницы)."""
    texts = []
    doc = fitz.open(path)
    for i, page in enumerate(doc):
        txt = page.get_text()
        texts.append((txt, i + 1))
    return texts

def _extract_docx(path: str) -> List[str]:
    """Возвращает список параграфов для DOCX.
    Приоритет: docx2txt → python-docx → zipfile fallback (без зависимостей).
    """
    paras: List[str] = []

    # 1) docx2txt
    if 'docx2txt' in globals() and HAS_DOCX2TXT:
        try:
            txt = docx2txt.process(path) or ""
            paras = [p.strip() for p in txt.split("\n\n") if p.strip()]
            parser_logger.info("DOCX распарсен через docx2txt")
            return paras
        except Exception as e:
            parser_logger.warning(f"docx2txt не справился ({e}), пробуем python-docx...")

    # 2) python-docx
    if 'DocxDocument' in globals() and HAS_PYTHON_DOCX:
        try:
            doc = DocxDocument(path)
            paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
            parser_logger.info("DOCX распарсен через python-docx")
            return paras
        except Exception as e:
            parser_logger.warning(f"python-docx не справился ({e}), пробуем zipfile fallback...")

    # 3) Fallback: читаем XML из DOCX-архива
    try:
        with zipfile.ZipFile(path, 'r') as z:
            xml = z.read('word/document.xml').decode('utf-8', errors='ignore')
        # вырезаем теги и приводим к «параграфам»
        text = re.sub(r"<[^>]+>", "\n", xml)           # теги → переводы строк
        text = re.sub(r"\n{2,}", "\n\n", text)         # нормализуем
        paras = [p.strip() for p in text.split("\n\n") if p.strip()]
        parser_logger.warning("DOCX распарсен через zipfile fallback (установите docx2txt или python-docx для лучшего качества).")
        return paras
    except Exception as e:
        # если совсем не получилось — лог и пустой список, чтобы не падать
        parser_logger.error(f"Не удалось распарсить DOCX ({e}). Файл будет пропущен.")
        return []

def _extract_txt(path: str) -> List[str]:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        data = f.read()
    return [p.strip() for p in data.split("\n\n") if p.strip()]

def _iter_documents_from_paths(file_paths: List[str]) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " "]
    )

    docs: List[Document] = []
    doc_id_counter = 0

    for path in file_paths:
        ext = os.path.splitext(path)[1].lower()
        parser_logger.info(f"Парсинг файла: {path}")

        if ext == ".pdf":
            pages = _extract_pdf(path)  # [(text, page_num), ...]
            for page_text, page_num in pages:
                # первичное разбиение по абзацам
                raw_paragraphs = [p.strip() for p in page_text.split("\n\n") if p.strip()]
                for para in raw_paragraphs:
                    if len(para.split()) > LONG_PARAGRAPH_WORDS:
                        # дополнительное разбиение
                        sub_docs = splitter.create_documents(
                            texts=[para],
                            metadatas=[{"file": os.path.basename(path), "page": page_num}]
                        )
                        for d in sub_docs:
                            d.metadata["id"] = doc_id_counter
                            doc_id_counter += 1
                        docs.extend(sub_docs)
                    else:
                        d = Document(
                            page_content=para,
                            metadata={"file": os.path.basename(path), "page": page_num, "id": doc_id_counter}
                        )
                        doc_id_counter += 1
                        docs.append(d)

        elif ext == ".docx":
            paras = _extract_docx(path)
            for para in paras:
                if len(para.split()) > LONG_PARAGRAPH_WORDS:
                    sub_docs = splitter.create_documents(
                        texts=[para],
                        metadatas=[{"file": os.path.basename(path), "page": 1}]
                    )
                    for d in sub_docs:
                        d.metadata["id"] = doc_id_counter
                        doc_id_counter += 1
                    docs.extend(sub_docs)
                else:
                    d = Document(
                        page_content=para,
                        metadata={"file": os.path.basename(path), "page": 1, "id": doc_id_counter}
                    )
                    doc_id_counter += 1
                    docs.append(d)

        elif ext == ".txt":
            paras = _extract_txt(path)
            for para in paras:
                if len(para.split()) > LONG_PARAGRAPH_WORDS:
                    sub_docs = splitter.create_documents(
                        texts=[para],
                        metadatas=[{"file": os.path.basename(path), "page": 1}]
                    )
                    for d in sub_docs:
                        d.metadata["id"] = doc_id_counter
                        doc_id_counter += 1
                    docs.extend(sub_docs)
                else:
                    d = Document(
                        page_content=para,
                        metadata={"file": os.path.basename(path), "page": 1, "id": doc_id_counter}
                    )
                    doc_id_counter += 1
                    docs.append(d)
        else:
            parser_logger.warning(f"Неподдерживаемое расширение: {path}")

    parser_logger.info(f"Итого фрагментов: {len(docs)}")
    return docs

# ----------------------------
# Агент-парсер
# ----------------------------
def parser_agent(state: GraphState):
    parser_logger.info("Старт парсинга")
    docs = _iter_documents_from_paths(state.file_paths)

    # здесь у вас уже есть цикл/логика по 413; добавим ретрай по 429
    max_tokens = MAX_EMBED_TOKENS_START
    vectorstore = None
    last_err = None

    while max_tokens >= MIN_EMBED_TOKENS:
        try:
            sanitized_docs = _sanitize_docs_for_embeddings(docs, max_tokens)
            texts = [d.page_content for d in sanitized_docs]
            metadatas = [d.metadata for d in sanitized_docs]

            # Ретраи на случай 429
            for attempt in range(EMBED_MAX_RETRIES):
                try:
                    vectorstore = FAISS.from_texts(texts=texts, embedding=embeddings, metadatas=metadatas)
                    break
                except Exception as e:
                    last_err = e
                    if _is_rate_limit_error(e):
                        delay = EMBED_BASE_DELAY * (2 ** attempt) + random.uniform(0, 0.4)
                        parser_logger.warning(f"429 при построении FAISS. Повтор через {delay:.1f}s "
                                              f"(попытка {attempt+1}/{EMBED_MAX_RETRIES})")
                        time.sleep(delay)
                        continue
                    else:
                        raise
            if vectorstore is not None:
                break  # успех
        except Exception as e:
            es = str(e)
            last_err = e
            if "413" in es or "Max tokens" in es:
                parser_logger.warning(f"413 (слишком длинный фрагмент). Уменьшаем лимит: {max_tokens} → {max_tokens - TOKEN_STEP}")
                max_tokens -= TOKEN_STEP
                continue
            else:
                parser_logger.error(f"Сбой построения FAISS: {e}")
                raise

    if vectorstore is None:
        raise RuntimeError(f"Не удалось построить FAISS: {last_err}")

    state.docs = sanitized_docs
    state.vectorstore = vectorstore
    parser_logger.info("Завершение парсинга и построения FAISS")
    return state

# ----------------------------
# Вспомогательные функции анализа
# ----------------------------
def _format_context_pack(docs: List[Document]) -> str:
    """Формирует компактный контекст-пакет с ID, файлом и страницей."""
    parts = []
    total_chars = 0
    for d in docs:
        sid = d.metadata.get("id")
        file = d.metadata.get("file")
        page = d.metadata.get("page")
        text = re.sub(r"\s+", " ", d.page_content).strip()
        text = text[:1000]  # ограничим размер одного блока
        block = f"[#{sid}] file={file}; page={page}; text: {text}"
        if total_chars + len(block) > MAX_CTX_CHARS:
            break
        parts.append(block)
        total_chars += len(block)
    return "\n".join(parts)

JSON_INSTRUCTIONS = """
Верни строго JSON без лишнего текста и форматирования. Структура:
{
  "answer": "<краткий ответ или 'not_found'>",
  "confidence": <число от 0 до 1>,
  "citations": [{"id": <int>, "file": "<str>", "page": <int>}],
  "quotes": ["<цитата 1>", "<цитата 2>"]
}
Если у тебя нет достаточных оснований ответить — используй "answer": "not_found".
Ответ строго основывай на контексте и обязательно укажи хотя бы один id из 'citations', если answer != "not_found".
"""

def _ask_llm_for_question(question: str, ctx_docs: List[Document]) -> Dict:
    """Один запрос LLM для одного вопроса по уже отобранным контекстам."""
    context_pack = _format_context_pack(ctx_docs)
    prompt = (
        "Ты аналитик ИТ-процессов. Ответь по контексту.\n\n"
        f"Вопрос: {question}\n\n"
        "Контекст:\n"
        f"{context_pack}\n\n"
        f"{JSON_INSTRUCTIONS}"
    )
    try:
        raw = llm.invoke(prompt, config={"timeout": 90}).content.strip()
        # Вырезаем JSON на случай обрамления
        m = re.search(r"\{.*\}", raw, flags=re.S)
        raw_json = m.group(0) if m else raw
        data = json.loads(raw_json)
        # Базовая валидация
        if not isinstance(data, dict) or "answer" not in data:
            raise ValueError("Bad JSON shape")
        return data
    except Exception as e:
        analysis_logger.error(f"Ошибка парсинга JSON для вопроса '{question}': {e}")
        return {"answer": "not_found", "confidence": 0.0, "citations": [], "quotes": []}

# ----------------------------
# Агент-анализатор
# ----------------------------
def analysis_agent(state: GraphState):
    analysis_logger.info("Старт анализа по вопросам (предэмбеддинг → поиск по вектору → LLM)")
    state.analysis_result = {q: [] for q in state.questions}

    if not state.vectorstore or not state.docs:
        analysis_logger.error("Нет векторного хранилища или документов.")
        return state

    # 1) Предэмбеддинг всех вопросов (последовательно, с бэкоффом и сериализацией)
    question_vecs: Dict[str, List[float]] = {}
    for q in state.questions:
        vec = embed_query_with_backoff(q)
        question_vecs[q] = vec
        analysis_logger.info(f"Эмбеддинг вопроса готов: '{q[:60]}...'")

    def _format_context_pack(docs: List[Document]) -> str:
        parts, total = [], 0
        for d in docs:
            sid  = d.metadata.get("id")
            file = d.metadata.get("file")
            page = d.metadata.get("page")
            text = re.sub(r"\s+", " ", d.page_content).strip()[:1000]
            block = f"[#{sid}] file={file}; page={page}; text: {text}"
            if total + len(block) > MAX_CTX_CHARS:
                break
            parts.append(block)
            total += len(block)
        return "\n".join(parts)

    def _ask_llm_for_question(question: str, ctx_docs: List[Document]) -> Dict:
        context_pack = _format_context_pack(ctx_docs)
        prompt = (
            "Ты аналитик ИТ-процессов. Ответь строго по контексту.\n\n"
            f"Вопрос: {question}\n\n"
            "Контекст:\n"
            f"{context_pack}\n\n"
            + JSON_INSTRUCTIONS
        )
        try:
            raw = llm_invoke_with_backoff(prompt, timeout=90)
            m = re.search(r"\{.*\}", raw, flags=re.S)
            raw_json = m.group(0) if m else raw
            data = json.loads(raw_json)
            if not isinstance(data, dict) or "answer" not in data:
                raise ValueError("Bad JSON shape")
            return data
        except Exception as e:
            analysis_logger.error(f"Ошибка парсинга JSON для '{question}': {e}")
            return {"answer": "not_found", "confidence": 0.0, "citations": [], "quotes": []}

    # 2) Параллельная обработка вопросов (векторный поиск не бьёт API)
    def process_one_question(q: str) -> Tuple[str, Dict]:
        q_vec = question_vecs[q]
        # Поиск по вектору — без вызова Embeddings API
        ctx_docs: List[Document] = state.vectorstore.similarity_search_by_vector(q_vec, k=TOP_K)
        res = _ask_llm_for_question(q, ctx_docs)

        # Подстрахуем citations (если LLM не вернул file/page)
        fixed_citations = []
        ids_to_doc = {d.metadata.get("id"): d for d in ctx_docs if "id" in d.metadata}
        for cit in res.get("citations", []) or []:
            try:
                cid = cit.get("id")
                d = ids_to_doc.get(cid)
                if d:
                    fixed_citations.append({"id": cid, "file": d.metadata.get("file"), "page": d.metadata.get("page")})
            except Exception:
                continue
        if not fixed_citations and res.get("answer") != "not_found" and ctx_docs:
            d0 = ctx_docs[0]
            fixed_citations = [{"id": d0.metadata.get("id"), "file": d0.metadata.get("file"), "page": d0.metadata.get("page")}]
        res["citations"] = fixed_citations
        return q, res

    from concurrent.futures import ThreadPoolExecutor, as_completed
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as ex:
        futures = [ex.submit(process_one_question, q) for q in state.questions]
        for fut in as_completed(futures):
            q, res = fut.result()
            analysis_logger.info(f"Вопрос: {q} | Ответ: {res.get('answer')} | conf={res.get('confidence')}")
            if res.get("answer") and res.get("answer") != "not_found":
                state.analysis_result[q].append(res)

    analysis_logger.info("Завершение анализа")
    return state


# ----------------------------
# Агент-отчёт
# ----------------------------
def report_agent(state: GraphState):
    report_logger.info("Старт построения отчёта")
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(True, 15)

    # Шрифты
    pdf.add_font("DejaVu", "", FONT_REGULAR)
    pdf.add_font("DejaVu", "B", FONT_BOLD)

    # Заголовок
    pdf.set_font("DejaVu", "B", 16)
    pdf.set_text_color(0, 102, 204)
    pdf.cell(0, 12, "Отчёт по анализу документа(ов)",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.ln(4)

    # Перечень файлов
    pdf.set_font("DejaVu", "", 11)
    pdf.set_text_color(60, 60, 60)
    pdf.set_x(pdf.l_margin)
    files_line = "Анализировались файлы: " + ", ".join([os.path.basename(p) for p in state.file_paths])
    pdf.multi_cell(0, 6, prepare_text_for_pdf(files_line),
                   new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(4)

    any_found = False
    for question, items in state.analysis_result.items():
        pdf.set_font("DejaVu", "B", 13)
        pdf.set_text_color(0, 0, 0)
        pdf.set_x(pdf.l_margin)
        pdf.cell(0, 9, f"Вопрос: {prepare_text_for_pdf(question)}",
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1)

        if items:
            any_found = True
            for item in items:
                ans   = prepare_text_for_pdf(item.get("answer", ""))
                conf  = item.get("confidence", 0) or 0
                quotes = item.get("quotes", []) or []
                cits   = item.get("citations", []) or []

                # Ответ
                pdf.set_font("DejaVu", "B", 12)
                pdf.set_text_color(0, 128, 0)
                pdf.set_x(pdf.l_margin)
                pdf.multi_cell(0, 7, f"Ответ: {ans} (уверенность: {conf:.2f})",
                               new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.ln(1)

                # Цитаты
                if quotes:
                    pdf.set_font("DejaVu", "", 11)
                    pdf.set_text_color(40, 40, 40)
                    for qtxt in quotes[:3]:
                        qtxt = prepare_text_for_pdf(qtxt)
                        pdf.set_x(pdf.l_margin)
                        pdf.multi_cell(0, 6, f'Цитата: "{qtxt}"',
                                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                    pdf.ln(1)

                # Источники
                if cits:
                    src_line = "Источники: " + "; ".join(
                        [f"{c.get('file')} стр.{c.get('page')} [#{c.get('id')}]"
                         for c in cits]
                    )
                    pdf.set_font("DejaVu", "", 11)
                    pdf.set_text_color(80, 80, 80)
                    pdf.set_x(pdf.l_margin)
                    pdf.multi_cell(0, 6, prepare_text_for_pdf(src_line),
                                   new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.ln(3)
        # else: ничего не выводим (только найденное)

        pdf.ln(2)

    if not any_found:
        pdf.set_font("DejaVu", "", 12)
        pdf.set_text_color(200, 0, 0)
        pdf.set_x(pdf.l_margin)
        pdf.cell(0, 8, "Ответы не найдены по заданным вопросам.",
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # Колонтитул
    pdf.set_y(-15)
    pdf.set_font("DejaVu", "", 10)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 10, f"Стр. {pdf.page_no()}", align="C")

    out_path = "analysis_report.pdf"
    pdf.output(out_path)
    report_logger.info(f"Отчёт сохранён: {out_path}")
    print(f"Отчёт сохранён: {out_path}")
    return state


# ----------------------------
# Коммуникатор / запуск графа
# ----------------------------
def _collect_files(input_path_or_list) -> List[str]:
    """Принимает строку (путь к папке/файлу) или список путей. Возвращает валидный список файлов."""
    if isinstance(input_path_or_list, list):
        files = input_path_or_list
    elif isinstance(input_path_or_list, str):
        if os.path.isdir(input_path_or_list):
            files = [
                os.path.join(input_path_or_list, f)
                for f in os.listdir(input_path_or_list)
                if f.lower().endswith((".pdf", ".docx", ".txt"))
            ]
        else:
            files = [input_path_or_list]
    else:
        raise ValueError("Передайте путь к папке/файлу или список путей.")
    files = [f for f in files if os.path.isfile(f)]
    if not files:
        raise FileNotFoundError("Не найдено ни одного входного файла.")
    return files

def communicate(inputs, questions: List[str]):
    comm_logger.info("Запуск графа")
    files = _collect_files(inputs)

    graph = StateGraph(GraphState)
    graph.add_node("parser", parser_agent)
    graph.add_node("analysis", analysis_agent)
    graph.add_node("report", report_agent)
    graph.set_entry_point("parser")
    graph.add_edge("parser", "analysis")
    graph.add_edge("analysis", "report")
    graph.add_edge("report", END)

    app = graph.compile()
    state = GraphState(file_paths=files, questions=questions)
    app.invoke(state)
    comm_logger.info("Граф завершён")

# ----------------------------
# Пример запуска
# ----------------------------
if __name__ == "__main__":
    # можно передать путь к папке с файлами:
    inputs = "docs"  # либо ["docs/a.pdf", "docs/b.docx", ...]
    questions = [
        "Артефакт по описанию технической архитектуры выполнен аккуратно?",
        "Артефакт по описанию технической архитектуры имеет отметки о согласовании с заказчиками и/или смежными подразделениями?",
        "Артефакт по описанию технической архитектуры имеет историю изменений?",
        "Процесс формирования артефакта по описанию технической архитектуры идентифицирован?",
        "Процесс формирования артефакта по описанию технической архитектуры исполняется регулярно по событию?",
        "Записи в истории изменений артефакта технической архитектуры трассируются  на вехи/проекты в компании",
        "Есть ли действующие ВНД / ОРД по формализации процесса по описанию технической архитектуры?",
        "На роли в процессе описания технической архитектуры определены соответствующие сотрудники?",
        "Назначения сотрудников на роли процесса закреплены в соответствующих ВНД/ОРД?",
        "Для процесса по описанию технической архитектуры установлены показатели RPI??"
    ]
    communicate(inputs, questions)
